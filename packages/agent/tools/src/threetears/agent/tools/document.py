"""Document reader -- parse documents into clean markdown.

Supports PDF, DOCX, XLSX, CSV, TXT, Markdown, and LaTeX formats.
Single entry point ``parse_document()`` dispatches by MIME type.
All sync parsers run via ``asyncio.to_thread()`` for non-blocking I/O.

Optional dependencies:
- ``pdf``: PyMuPDF (fitz) for PDF parsing
- ``docx``: python-docx for DOCX parsing
- ``xlsx``: openpyxl for XLSX parsing
- ``ocr``: pytesseract + pdf2image for scanned PDF OCR
"""

from __future__ import annotations

import asyncio
import base64
import csv
import io
import mimetypes
import re
from dataclasses import dataclass, field
from typing import Any

from langchain_core.tools import BaseTool, tool
from pydantic import BaseModel, Field

from threetears.agent.tools.base_tool import MCPToolDefinition, TearsTool, ToolResult
from threetears.observe import get_logger, traced

__all__ = [
    "DocumentResult",
    "DocumentSection",
    "OcrConfig",
    "ParseDocumentInput",
    "ParseDocumentTool",
    "create_parse_document_tool",
    "detect_mime_from_filename",
    "parse_document",
]

log = get_logger(__name__)


# -- Data types ---------------------------------------------------------------


@dataclass
class OcrConfig:
    """Configuration for OCR fallback on scanned PDF pages."""

    enabled: bool = False
    language: str = "eng"


@dataclass
class DocumentSection:
    """A section within a parsed document."""

    heading: str | None
    content: str
    page_number: int | None
    level: int  # heading level 1-6, 0 for body


@dataclass
class DocumentResult:
    """Result of parsing a document into markdown."""

    text: str  # Full document as clean markdown
    title: str | None
    page_count: int | None
    word_count: int
    was_ocr: bool
    metadata: dict[str, Any] = field(default_factory=dict)
    sections: list[DocumentSection] = field(default_factory=list)


# -- MIME type dispatch -------------------------------------------------------

_MIME_PARSERS: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "text/csv": "csv",
    "text/plain": "txt",
    "text/markdown": "markdown",
    "application/x-tex": "latex",
    "application/x-latex": "latex",
}


def detect_mime_from_filename(filename: str) -> str | None:
    """Detect MIME type from filename extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    ext_map = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "csv": "text/csv",
        "txt": "text/plain",
        "md": "text/markdown",
        "tex": "application/x-tex",
        "latex": "application/x-latex",
    }
    return ext_map.get(ext) or mimetypes.guess_type(filename)[0]


@traced()
async def parse_document(
    data: bytes,
    mime_type: str,
    filename: str | None = None,
    ocr_config: OcrConfig | None = None,
) -> DocumentResult:
    """Parse document bytes into markdown.

    Dispatches to format-specific parsers based on MIME type.
    Falls back to filename extension detection if MIME type is unknown.
    """
    parser_key = _MIME_PARSERS.get(mime_type)

    # Fallback to filename extension
    if parser_key is None and filename:
        detected = detect_mime_from_filename(filename)
        if detected:
            parser_key = _MIME_PARSERS.get(detected)

    if parser_key is None:
        return DocumentResult(
            text=f"[Unsupported document type: {mime_type}]",
            title=filename,
            page_count=None,
            word_count=0,
            was_ocr=False,
        )

    parsers = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "xlsx": _parse_xlsx,
        "csv": _parse_csv,
        "txt": _parse_txt,
        "markdown": _parse_markdown,
        "latex": _parse_latex,
    }

    parser_fn = parsers[parser_key]
    ocr = ocr_config or OcrConfig()
    return await asyncio.to_thread(parser_fn, data, filename, ocr)


# -- PDF parser ---------------------------------------------------------------


def _parse_pdf(
    data: bytes,
    filename: str | None = None,
    ocr: OcrConfig = OcrConfig(),
) -> DocumentResult:
    """Parse PDF using PyMuPDF with OCR fallback for scanned pages."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=data, filetype="pdf")
        try:
            sections: list[DocumentSection] = []
            full_parts: list[str] = []
            was_ocr = False
            title = doc.metadata.get("title") or filename

            metadata = {}
            for key in ("author", "subject", "keywords", "creator", "producer"):
                val = doc.metadata.get(key)
                if val:
                    metadata[key] = val

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text")

                # OCR fallback for pages with very little text
                if len(text.strip()) < 50 and ocr.enabled:
                    ocr_text = _ocr_page(data, page_num, ocr.language)
                    if ocr_text:
                        text = ocr_text
                        was_ocr = True

                # Try to extract tables
                tables_md = _extract_pdf_tables(page)

                # Font-size heuristic for headings
                page_sections = _extract_pdf_headings(page, text, page_num + 1)
                if page_sections:
                    sections.extend(page_sections)
                else:
                    sections.append(
                        DocumentSection(
                            heading=None,
                            content=text.strip(),
                            page_number=page_num + 1,
                            level=0,
                        )
                    )

                page_text = text.strip()
                if tables_md:
                    page_text += "\n\n" + tables_md
                if page_text:
                    full_parts.append(page_text)

            page_count = len(doc)
        finally:
            doc.close()
        full_text = "\n\n---\n\n".join(full_parts)
        word_count = len(full_text.split())

        return DocumentResult(
            text=full_text,
            title=title if title else None,
            page_count=page_count,
            word_count=word_count,
            was_ocr=was_ocr,
            metadata=metadata,
            sections=sections,
        )

    except Exception as exc:
        log.error(
            "PDF parsing failed",
            extra={"extra_data": {"error": str(exc)}},
        )
        return DocumentResult(
            text=f"[Parsing failed: {exc}]",
            title=filename,
            page_count=None,
            word_count=0,
            was_ocr=False,
        )


def _extract_pdf_tables(page: Any) -> str:
    """Extract tables from a PDF page as markdown."""
    try:
        tables = page.find_tables()
        if not tables or not tables.tables:
            return ""

        parts = []
        for table in tables.tables:
            rows = table.extract()
            if not rows:
                continue
            # Build markdown table
            header = rows[0]
            header_line = "| " + " | ".join(str(c) if c else "" for c in header) + " |"
            sep_line = "| " + " | ".join("---" for _ in header) + " |"
            body_lines = []
            for row in rows[1:]:
                cells = [str(c) if c else "" for c in row]
                body_lines.append("| " + " | ".join(cells) + " |")
            parts.append("\n".join([header_line, sep_line] + body_lines))

        return "\n\n".join(parts)
    except Exception:
        return ""


def _extract_pdf_headings(
    page: Any,
    text: str,
    page_number: int,
) -> list[DocumentSection]:
    """Use font-size heuristic to identify headings on a PDF page."""
    try:
        blocks = page.get_text("dict")["blocks"]
        if not blocks:
            return []

        font_sizes: list[float] = []
        for block in blocks:
            if block.get("type") == 0:
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        font_sizes.append(span.get("size", 12))

        if not font_sizes:
            return []

        median_size = sorted(font_sizes)[len(font_sizes) // 2]

        sections: list[DocumentSection] = []
        current_heading: str | None = None
        current_level = 0
        current_text: list[str] = []

        for line in text.split("\n"):
            line_text = line.strip()
            if not line_text:
                continue

            is_heading = False
            for block in blocks:
                if block.get("type") == 0:
                    for bline in block.get("lines", []):
                        for span in bline.get("spans", []):
                            if span.get("text", "").strip() == line_text and span.get("size", 12) > median_size * 1.2:
                                is_heading = True

            if is_heading:
                if current_text or current_heading:
                    sections.append(
                        DocumentSection(
                            heading=current_heading,
                            content="\n".join(current_text).strip(),
                            page_number=page_number,
                            level=current_level,
                        )
                    )
                    current_text = []

                if any(
                    span.get("text", "").strip() == line_text and span.get("size", 12) > median_size * 1.5
                    for block in blocks
                    if block.get("type") == 0
                    for bline in block.get("lines", [])
                    for span in bline.get("spans", [])
                ):
                    current_level = 1
                else:
                    current_level = 3
                current_heading = line_text
            else:
                current_text.append(line_text)

        # Flush remaining
        if current_text or current_heading:
            sections.append(
                DocumentSection(
                    heading=current_heading,
                    content="\n".join(current_text).strip(),
                    page_number=page_number,
                    level=current_level,
                )
            )

        return sections

    except Exception:
        return []


def _ocr_page(pdf_data: bytes, page_num: int, language: str) -> str | None:
    """OCR a single PDF page using pytesseract + pdf2image."""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract

        images = convert_from_bytes(
            pdf_data,
            first_page=page_num + 1,
            last_page=page_num + 1,
            dpi=300,
        )
        if not images:
            return None

        text = pytesseract.image_to_string(images[0], lang=language)
        return text.strip() if text.strip() else None

    except Exception as exc:
        log.warning(
            "OCR failed for page",
            extra={"extra_data": {"page": page_num, "error": str(exc)}},
        )
        return None


# -- DOCX parser --------------------------------------------------------------


def _parse_docx(
    data: bytes,
    filename: str | None = None,
    ocr: OcrConfig = OcrConfig(),
) -> DocumentResult:
    """Parse DOCX using python-docx."""
    try:
        import io

        from docx import Document

        doc = Document(io.BytesIO(data))
        sections: list[DocumentSection] = []
        full_parts: list[str] = []
        title = None

        # Extract core properties
        metadata: dict[str, Any] = {}
        try:
            props = doc.core_properties
            if props.title:
                title = props.title
            for attr in ("author", "subject", "keywords", "category"):
                val = getattr(props, attr, None)
                if val:
                    metadata[attr] = val
        except Exception:
            pass

        current_heading: str | None = None
        current_level = 0
        current_text: list[str] = []

        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower() if para.style else ""

            # Detect headings
            if style_name.startswith("heading"):
                # Flush current section
                if current_text or current_heading:
                    sections.append(
                        DocumentSection(
                            heading=current_heading,
                            content="\n".join(current_text).strip(),
                            page_number=None,
                            level=current_level,
                        )
                    )
                    full_parts.extend(current_text)
                    current_text = []

                # Parse heading level
                try:
                    level = int(style_name.replace("heading", "").strip())
                except ValueError:
                    level = 1

                current_heading = para.text.strip()
                current_level = level
                md_heading = "#" * level + " " + current_heading
                full_parts.append(md_heading)
            else:
                # Process inline formatting
                line = _docx_para_to_markdown(para)
                if line:
                    current_text.append(line)

        # Flush remaining
        if current_text or current_heading:
            sections.append(
                DocumentSection(
                    heading=current_heading,
                    content="\n".join(current_text).strip(),
                    page_number=None,
                    level=current_level,
                )
            )
            full_parts.extend(current_text)

        # Extract tables
        for table in doc.tables:
            table_md = _docx_table_to_markdown(table)
            if table_md:
                full_parts.append(table_md)

        full_text = "\n\n".join(p for p in full_parts if p.strip())
        word_count = len(full_text.split())

        return DocumentResult(
            text=full_text,
            title=title or filename,
            page_count=None,
            word_count=word_count,
            was_ocr=False,
            metadata=metadata,
            sections=sections,
        )

    except Exception as exc:
        log.error(
            "DOCX parsing failed",
            extra={"extra_data": {"error": str(exc)}},
        )
        return DocumentResult(
            text=f"[Parsing failed: {exc}]",
            title=filename,
            page_count=None,
            word_count=0,
            was_ocr=False,
        )


def _docx_para_to_markdown(para: Any) -> str:
    """Convert a DOCX paragraph to markdown with inline formatting."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts)


def _docx_table_to_markdown(table: Any) -> str:
    """Convert a DOCX table to markdown."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")

    if len(rows) < 1:
        return ""

    # Add separator after header
    header_cells = len(table.rows[0].cells)
    sep = "| " + " | ".join("---" for _ in range(header_cells)) + " |"
    return rows[0] + "\n" + sep + "\n" + "\n".join(rows[1:])


# -- XLSX parser --------------------------------------------------------------


def _parse_xlsx(
    data: bytes,
    filename: str | None = None,
    ocr: OcrConfig = OcrConfig(),
) -> DocumentResult:
    """Parse XLSX using openpyxl."""
    try:
        import io

        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
        try:
            sections: list[DocumentSection] = []
            full_parts: list[str] = []

            metadata: dict[str, Any] = {"sheet_count": len(wb.sheetnames)}

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                heading = f"## {sheet_name}"
                full_parts.append(heading)

                # Read all rows
                all_rows: list[list[str]] = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    # Skip completely empty rows
                    if any(c.strip() for c in cells):
                        all_rows.append(cells)

                if not all_rows:
                    sections.append(
                        DocumentSection(
                            heading=sheet_name,
                            content="(empty sheet)",
                            page_number=None,
                            level=2,
                        )
                    )
                    continue

                # Build markdown table
                header = all_rows[0]
                header_line = "| " + " | ".join(header) + " |"
                sep_line = "| " + " | ".join("---" for _ in header) + " |"
                body_lines = []
                for row in all_rows[1:]:
                    # Pad or trim to header width
                    padded = row + [""] * (len(header) - len(row))
                    body_lines.append("| " + " | ".join(padded[: len(header)]) + " |")

                table_md = "\n".join([header_line, sep_line] + body_lines)
                full_parts.append(table_md)

                sections.append(
                    DocumentSection(
                        heading=sheet_name,
                        content=table_md,
                        page_number=None,
                        level=2,
                    )
                )
        finally:
            wb.close()
        full_text = "\n\n".join(full_parts)
        word_count = len(full_text.split())

        return DocumentResult(
            text=full_text,
            title=filename,
            page_count=None,
            word_count=word_count,
            was_ocr=False,
            metadata=metadata,
            sections=sections,
        )

    except Exception as exc:
        log.error(
            "XLSX parsing failed",
            extra={"extra_data": {"error": str(exc)}},
        )
        return DocumentResult(
            text=f"[Parsing failed: {exc}]",
            title=filename,
            page_count=None,
            word_count=0,
            was_ocr=False,
        )


# -- CSV parser ---------------------------------------------------------------


def _parse_csv(
    data: bytes,
    filename: str | None = None,
    ocr: OcrConfig = OcrConfig(),
) -> DocumentResult:
    """Parse CSV into a single markdown table using the stdlib ``csv`` module.

    Already tabular -- unlike XLSX/DOCX, there's no source-format table
    structure to walk, just rows to read (``csv.reader`` handles RFC 4180
    quoting -- embedded commas/newlines inside quoted fields -- correctly,
    unlike a naive ``line.split(",")``). Mirrors ``_parse_xlsx``'s single-
    table shape: first non-empty row is the header, every row after is
    padded/trimmed to the header's width so a ragged CSV still produces a
    well-formed table.
    """
    try:
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("latin-1")

        all_rows: list[list[str]] = []
        for row in csv.reader(io.StringIO(text)):
            if any(cell.strip() for cell in row):
                all_rows.append(row)

        if not all_rows:
            return DocumentResult(
                text="(empty file)",
                title=filename,
                page_count=None,
                word_count=0,
                was_ocr=False,
            )

        header = all_rows[0]
        header_line = "| " + " | ".join(header) + " |"
        sep_line = "| " + " | ".join("---" for _ in header) + " |"
        body_lines = []
        for row in all_rows[1:]:
            padded = row + [""] * (len(header) - len(row))
            body_lines.append("| " + " | ".join(padded[: len(header)]) + " |")

        full_text = "\n".join([header_line, sep_line] + body_lines)
        word_count = len(full_text.split())

        sections = [
            DocumentSection(
                heading=None,
                content=full_text,
                page_number=None,
                level=0,
            )
        ]

        return DocumentResult(
            text=full_text,
            title=filename,
            page_count=None,
            word_count=word_count,
            was_ocr=False,
            sections=sections,
        )

    except Exception as exc:
        log.error(
            "CSV parsing failed",
            extra={"extra_data": {"error": str(exc)}},
        )
        return DocumentResult(
            text=f"[Parsing failed: {exc}]",
            title=filename,
            page_count=None,
            word_count=0,
            was_ocr=False,
        )


# -- TXT parser ---------------------------------------------------------------


def _parse_txt(
    data: bytes,
    filename: str | None = None,
    ocr: OcrConfig = OcrConfig(),
) -> DocumentResult:
    """Parse plain text with encoding fallback."""
    try:
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("latin-1")

        word_count = len(text.split())

        sections = [
            DocumentSection(
                heading=None,
                content=text,
                page_number=None,
                level=0,
            )
        ]

        return DocumentResult(
            text=text,
            title=filename,
            page_count=None,
            word_count=word_count,
            was_ocr=False,
            sections=sections,
        )

    except Exception as exc:
        return DocumentResult(
            text=f"[Parsing failed: {exc}]",
            title=filename,
            page_count=None,
            word_count=0,
            was_ocr=False,
        )


# -- Markdown parser ----------------------------------------------------------


def _parse_markdown(
    data: bytes,
    filename: str | None = None,
    ocr: OcrConfig = OcrConfig(),
) -> DocumentResult:
    """Parse markdown -- direct passthrough since it's already markdown."""
    try:
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("latin-1")

        word_count = len(text.split())

        # Extract sections from markdown headings
        sections: list[DocumentSection] = []
        current_heading: str | None = None
        current_level = 0
        current_text: list[str] = []

        for line in text.split("\n"):
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading_match:
                if current_text or current_heading:
                    sections.append(
                        DocumentSection(
                            heading=current_heading,
                            content="\n".join(current_text).strip(),
                            page_number=None,
                            level=current_level,
                        )
                    )
                    current_text = []

                current_level = len(heading_match.group(1))
                current_heading = heading_match.group(2).strip()
            else:
                current_text.append(line)

        if current_text or current_heading:
            sections.append(
                DocumentSection(
                    heading=current_heading,
                    content="\n".join(current_text).strip(),
                    page_number=None,
                    level=current_level,
                )
            )

        # Try to extract title from first heading
        title = filename
        if sections and sections[0].heading:
            title = sections[0].heading

        return DocumentResult(
            text=text,
            title=title,
            page_count=None,
            word_count=word_count,
            was_ocr=False,
            sections=sections,
        )

    except Exception as exc:
        return DocumentResult(
            text=f"[Parsing failed: {exc}]",
            title=filename,
            page_count=None,
            word_count=0,
            was_ocr=False,
        )


# -- LaTeX parser -------------------------------------------------------------


def _parse_latex(
    data: bytes,
    filename: str | None = None,
    ocr: OcrConfig = OcrConfig(),
) -> DocumentResult:
    """Parse LaTeX into markdown using regex transformations."""
    try:
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("latin-1")

        # Extract title before stripping preamble
        title_match = re.search(r"\\title\{([^}]+)\}", text)
        title = title_match.group(1) if title_match else filename

        # Strip preamble
        doc_start = re.search(r"\\begin\{document\}", text)
        doc_end = re.search(r"\\end\{document\}", text)
        if doc_start:
            text = text[doc_start.end() :]
        if doc_end:
            text = text[: doc_end.start()]

        # Remove common preamble commands
        text = re.sub(r"\\(usepackage|documentclass|author|date|maketitle)\{[^}]*\}", "", text)
        text = re.sub(r"\\(usepackage|documentclass)\[[^\]]*\]\{[^}]*\}", "", text)

        # Section headings
        text = re.sub(r"\\section\*?\{([^}]+)\}", r"# \1", text)
        text = re.sub(r"\\subsection\*?\{([^}]+)\}", r"## \1", text)
        text = re.sub(r"\\subsubsection\*?\{([^}]+)\}", r"### \1", text)
        text = re.sub(r"\\paragraph\*?\{([^}]+)\}", r"#### \1", text)

        # Inline formatting
        text = re.sub(r"\\textbf\{([^}]+)\}", r"**\1**", text)
        text = re.sub(r"\\textit\{([^}]+)\}", r"*\1*", text)
        text = re.sub(r"\\emph\{([^}]+)\}", r"*\1*", text)
        text = re.sub(r"\\texttt\{([^}]+)\}", r"`\1`", text)
        text = re.sub(r"\\underline\{([^}]+)\}", r"\1", text)

        # Code blocks
        text = re.sub(
            r"\\begin\{verbatim\}(.*?)\\end\{verbatim\}",
            r"```\n\1\n```",
            text,
            flags=re.DOTALL,
        )
        text = re.sub(
            r"\\begin\{lstlisting\}(.*?)\\end\{lstlisting\}",
            r"```\n\1\n```",
            text,
            flags=re.DOTALL,
        )

        # Lists
        text = re.sub(r"\\begin\{(itemize|enumerate)\}", "", text)
        text = re.sub(r"\\end\{(itemize|enumerate)\}", "", text)
        text = re.sub(r"\\item\s*", "- ", text)

        # Remove remaining LaTeX commands (simple ones)
        text = re.sub(r"\\[a-zA-Z]+\{([^}]*)\}", r"\1", text)

        # Clean up
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()

        word_count = len(text.split())

        # Extract sections
        sections: list[DocumentSection] = []
        current_heading: str | None = None
        current_level = 0
        current_text: list[str] = []

        for line in text.split("\n"):
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading_match:
                if current_text or current_heading:
                    sections.append(
                        DocumentSection(
                            heading=current_heading,
                            content="\n".join(current_text).strip(),
                            page_number=None,
                            level=current_level,
                        )
                    )
                    current_text = []
                current_level = len(heading_match.group(1))
                current_heading = heading_match.group(2).strip()
            else:
                current_text.append(line)

        if current_text or current_heading:
            sections.append(
                DocumentSection(
                    heading=current_heading,
                    content="\n".join(current_text).strip(),
                    page_number=None,
                    level=current_level,
                )
            )

        return DocumentResult(
            text=text,
            title=title,
            page_count=None,
            word_count=word_count,
            was_ocr=False,
            sections=sections,
        )

    except Exception as exc:
        return DocumentResult(
            text=f"[Parsing failed: {exc}]",
            title=filename,
            page_count=None,
            word_count=0,
            was_ocr=False,
        )


# -- parse_document tool factory ----------------------------------------------


class ParseDocumentInput(BaseModel):
    """Input schema for the parse_document tool."""

    content_base64: str = Field(description="Base64-encoded file content")
    filename: str = Field(
        description="Original filename with extension, used for format detection",
    )


_PARSE_DOCUMENT_MAX_CHARS = 15_000
_PARSE_DOCUMENT_MAX_BYTES = 50 * 1024 * 1024  # 50 MB


def _tool_error(tool_name: str, action: str, error: str) -> str:
    return f"[TOOL ERROR] {tool_name}: {action} failed — {error}"


def create_parse_document_tool(
    config: dict[str, Any],
    description: str,
    ocr_config: OcrConfig | None = None,
) -> BaseTool:
    """Create a LangChain parse_document tool.

    Decodes base64 content, detects MIME type from filename, parses the
    document, and returns clean markdown text (truncated to 15K chars).
    """
    ocr = ocr_config or OcrConfig()

    @tool("parse_document", args_schema=ParseDocumentInput)
    async def parse_document_tool(content_base64: str, filename: str) -> str:
        """Parse binary document content into clean markdown text."""
        try:
            data = base64.b64decode(content_base64)
        except Exception as exc:
            return _tool_error("parse_document", "decode", f"Invalid base64: {exc}")

        if len(data) > _PARSE_DOCUMENT_MAX_BYTES:
            return _tool_error(
                "parse_document",
                "decode",
                f"Decoded content exceeds maximum size ({len(data)} > {_PARSE_DOCUMENT_MAX_BYTES} bytes)",
            )

        mime_type = detect_mime_from_filename(filename)
        if mime_type is None:
            return _tool_error(
                "parse_document",
                "detect format",
                f"Cannot determine format from filename '{filename}'. Supported: .pdf, .docx, .xlsx, .csv, .txt, .md, .tex",
            )

        try:
            result = await parse_document(data, mime_type, filename, ocr_config=ocr)
        except Exception as exc:
            return _tool_error("parse_document", "parse", str(exc))

        text = result.text
        if len(text) > _PARSE_DOCUMENT_MAX_CHARS:
            text = text[:_PARSE_DOCUMENT_MAX_CHARS] + "\n\n[Content truncated]"

        # Build response with metadata
        parts = []
        if result.title:
            parts.append(f"**Title:** {result.title}")
        if result.page_count:
            parts.append(f"**Pages:** {result.page_count}")
        parts.append(f"**Words:** {result.word_count}")
        if result.was_ocr:
            parts.append("**Note:** OCR was used for some pages")
        parts.append("")
        parts.append(text)

        return "\n".join(parts)

    parse_document_tool.description = description
    return parse_document_tool


class ParseDocumentTool(TearsTool):
    """TearsTool wrapper for document parsing into clean markdown.

    parses binary document content (PDF, DOCX, XLSX, CSV, TXT, Markdown,
    LaTeX) into markdown text. accepts base64-encoded content and
    detects format from filename extension. optionally supports OCR
    for scanned PDF pages.
    """

    _INPUT_SCHEMA: dict[str, Any] = {
        "type": "object",
        "properties": {
            "content_base64": {
                "type": "string",
                "description": "base64-encoded file content",
            },
            "filename": {
                "type": "string",
                "description": "original filename with extension, used for format detection",
            },
        },
        "required": ["content_base64", "filename"],
    }

    def __init__(self, ocr_config: OcrConfig | None = None) -> None:
        """initialize parse document tool with optional OCR configuration.

        :param ocr_config: OCR configuration for scanned PDF fallback
        :ptype ocr_config: OcrConfig | None
        """
        self._ocr_config = ocr_config

    async def execute(self, **kwargs: Any) -> ToolResult:
        """parse document content into markdown.

        :param kwargs: must include 'content_base64' and 'filename' keys
        :ptype kwargs: Any
        :return: result containing parsed markdown or error
        :rtype: ToolResult
        """
        content_base64 = kwargs.get("content_base64", "")
        filename = kwargs.get("filename", "")
        lc_tool = create_parse_document_tool(
            config={},
            description="parse document",
            ocr_config=self._ocr_config,
        )
        content = await lc_tool.ainvoke(
            {"content_base64": content_base64, "filename": filename},
        )
        success = not content.startswith("[TOOL ERROR]")
        result = ToolResult(
            success=success,
            content=content,
            error=content if not success else None,
        )
        return result

    def mcp_schema(self) -> MCPToolDefinition:
        """return MCP-compatible tool definition for parse document.

        :return: tool definition with name, version, description, input schema
        :rtype: MCPToolDefinition
        """
        result = MCPToolDefinition(
            name=self.mcp_name(),
            version=self.mcp_version(),
            description="parse binary document content into clean markdown text",
            input_schema=self._INPUT_SCHEMA,
        )
        return result

    def mcp_name(self) -> str:
        """return namespaced tool name.

        :return: namespaced tool name
        :rtype: str
        """
        return "threetears.parse_document"

    def mcp_version(self) -> str:
        """return tool version.

        :return: version string
        :rtype: str
        """
        return "1.0"
